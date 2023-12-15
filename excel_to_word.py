from openpyxl import load_workbook  # openpyxl에서 load_workbook 함수 임포트
from docx import Document  # docx에서 Document 클래스 임포트

# 엑셀 파일 로드
workbook = load_workbook('test.xlsx')  # 엑셀 파일을 로드
sheet = workbook.active  # 활성화된 시트를 가져옴

# 워드 파일 생성
doc = Document()  # 새 워드 파일 생성
table = doc.add_table(rows=1, cols=6)  # 워드 파일에 표 추가 (행: 1, 열: 6)

# 헤더에 열 이름 추가
hdr_cells = table.rows[0].cells  # 표의 첫 번째 행을 헤더 셀로 지정
hdr_cells[0].text = '1번째 열'  # 첫 번째 열에 '1번째 열' 추가
hdr_cells[1].text = '2번째 열'  # 두 번째 열에 '2번째 열' 추가
hdr_cells[2].text = '3번째 열'  # 세 번째 열에 '3번째 열' 추가
hdr_cells[3].text = '4번째 열'  # 네 번째 열에 '4번째 열' 추가
hdr_cells[4].text = '5번째 열'  # 다섯 번째 열에 '5번째 열' 추가
hdr_cells[5].text = '6번째 열'  # 여섯 번째 열에 '6번째 열' 추가

# 엑셀에서 데이터를 읽어와서 워드 파일에 추가 (두 번째 행부터)
for row in sheet.iter_rows(min_row=1, values_only=True):
    # 엑셀 데이터의 각 행에 대해 이름, 변경 전, 변경 후 값을 가져옴
    name, before, after = row[0], row[1], row[2]

    # 표에 새로운 행 추가하여 데이터 추가
    row_cells = table.add_row().cells  # 표에 새로운 행 추가

    # 열마다 데이터 배치
    row_cells[0].text = str(name)  # 1번째 열에 이름 데이터 추가
    row_cells[2].text = str(before)  # 3번째 열에 변경 전 데이터 추가
    row_cells[3].text = str(after)  # 4번째 열에 변경 후 데이터 추가

    # 2번째 열을 2개의 행으로 분할하여 데이터 추가
    cell_2 = row_cells[1]
    cell_2.text = '역삼'  # 위쪽 칸에 '역삼' 추가
    cell_2.add_paragraph('분당')  # 아래쪽 칸에 '분당' 추가

    # 5번째 열에는 4번째 열에서 3번째 열 데이터를 뺀 값 추가
    result_value = row[2] - row[1]
    row_cells[4].text = str(result_value)  # 5번째 열에 결과값 추가

# 워드 파일 저장
doc.save('result_word_file_updated.docx')  # 완성된 워드 파일 저장
