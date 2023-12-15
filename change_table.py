from docx import Document
from copy import deepcopy

def find_table_start_position(doc, table_text):
    # 원하는 텍스트가 포함된 표를 찾아 해당 위치 반환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if table_text in cell.text:
                    return table._element
    return None

def insert_table_at_position(existing_doc, new_doc, table_text):
    # 기존 파일에서 표 삽입할 위치 찾기
    position = find_table_start_position(existing_doc, table_text)

    # 새로운 표 복사하여 삽입
    if position:
        new_table = deepcopy(new_doc.tables[0])  # 새 표의 첫 번째 표를 복사
        position.addnext(new_table._element)
        return True
    return False

# 기존의 워드 파일 열기
existing_doc = Document('양식파일.docx')

# 새롭게 생성한 워드 파일 열기
new_doc = Document('result_word_file_updated.docx')

# 기존 파일에서 삽입할 표를 위치를 찾을 기준 텍스트 지정
table_text = "표가 시작되는 특정한 문구"  # 기존 파일에서 표가 시작하는 위치를 대표하는 텍스트

# 기존 워드 파일에서 원하는 위치에 새로운 표 삽입
inserted = insert_table_at_position(existing_doc, new_doc, table_text)

# 변경된 기존 워드 파일 저장
if inserted:
    existing_doc.save('updated_existing_word_file.docx')
    print("표가 성공적으로 삽입되었습니다.")
else:
    print("표 삽입에 실패했습니다.")
